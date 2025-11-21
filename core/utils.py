# Utilitários para Otimização de Performance
from django.db.models import Prefetch, Count, Q
from django.core.cache import cache
from django.conf import settings
from core.models import Turma, Aluno, LancamentoDeNota, Professor, Competencia
import logging

logger = logging.getLogger(__name__)

class QueryOptimizer:
    """
    Classe para otimizar queries do banco de dados
    """
    
    @staticmethod
    def get_turmas_com_progresso_otimizado(professor=None):
        """
        Versão otimizada para buscar turmas com estatísticas de progresso
        Reduz significativamente o número de queries ao banco
        """
        # Query base com select_related e prefetch_related
        queryset = Turma.objects.select_related(
            'tipo_turma',
            'professor_responsavel__user'
        ).prefetch_related(
            'competencias',
            Prefetch(
                'alunos',
                queryset=Aluno.objects.only('id', 'nome_completo', 'turma_id')
            ),
            Prefetch(
                'alunos__lancamentos_de_nota',
                queryset=LancamentoDeNota.objects.select_related('competencia')
            )
        )
        
        # Filtrar por professor se especificado
        if professor:
            queryset = queryset.filter(professor_responsavel=professor)
        
        # Anotar com contagens
        queryset = queryset.annotate(
            total_alunos=Count('alunos', distinct=True),
            total_notas_lancadas=Count('alunos__lancamentos_de_nota', distinct=True)
        )
        
        return queryset
    
    @staticmethod
    def get_dashboard_admin_otimizado():
        """
        Query otimizada para o dashboard administrativo
        """
        cache_key = 'dashboard_admin_data'
        cached_data = cache.get(cache_key)
        
        if cached_data and not settings.DEBUG:
            return cached_data
        
        # Query principal otimizada
        turmas = Turma.objects.select_related(
            'tipo_turma',
            'professor_responsavel__user'
        ).prefetch_related(
            'competencias'
        ).annotate(
            total_alunos=Count('alunos', distinct=True),
            total_notas=Count('alunos__lancamentos_de_nota', distinct=True)
        )
        
        data = {
            'turmas': list(turmas),
            'totais': {
                'turmas': turmas.count(),
                'alunos': sum(getattr(t, 'total_alunos', 0) for t in turmas),
                'notas': sum(getattr(t, 'total_notas', 0) for t in turmas)
            }
        }
        
        # Cache por 5 minutos
        cache.set(cache_key, data, 300)
        return data

class CacheManager:
    """
    Gerenciador de cache para dados frequentemente acessados
    """
    
    @staticmethod
    def get_or_set_cache(key, callable_func, timeout=300):
        """
        Busca no cache ou executa função e armazena resultado
        """
        data = cache.get(key)
        if data is None:
            data = callable_func()
            cache.set(key, data, timeout)
        return data
    
    @staticmethod
    def invalidate_turma_cache(turma_id):
        """
        Invalida cache relacionado a uma turma específica
        """
        keys_to_delete = [
            f'turma_progresso_{turma_id}',
            f'turma_detalhes_{turma_id}',
            'dashboard_admin_data',
            'dashboard_analytics_data'
        ]
        cache.delete_many(keys_to_delete)
    
    @staticmethod
    def invalidate_professor_cache(professor_id):
        """
        Invalida cache relacionado a um professor específico
        """
        keys_to_delete = [
            f'professor_dashboard_{professor_id}',
            'dashboard_admin_data'
        ]
        cache.delete_many(keys_to_delete)

class DataValidator:
    """
    Validações de dados e integridade
    """
    
    @staticmethod
    def validate_nota_valor(valor, tipo_nota):
        """
        Valida se o valor da nota está correto para o tipo
        """
        if tipo_nota == 'NUM':
            try:
                num_valor = float(valor)
                if 0 <= num_valor <= 100:
                    return True, num_valor
                else:
                    return False, "Nota numérica deve estar entre 0 e 100"
            except ValueError:
                return False, "Valor deve ser um número"
        
        elif tipo_nota == 'ABC':
            if valor.upper() in ['A', 'B', 'C', 'D']:
                return True, valor.upper()
            else:
                return False, "Nota conceitual deve ser A, B, C ou D"
        
        return False, "Tipo de nota inválido"
    
    @staticmethod
    def validate_import_data(df):
        """
        Valida dados de importação em lote
        """
        errors = []
        warnings = []
        
        # Verificar colunas obrigatórias
        required_columns = ['nome_completo']
        for col in required_columns:
            if col not in df.columns:
                errors.append(f"Coluna obrigatória '{col}' não encontrada")
        
        # Verificar dados vazios
        if df['nome_completo'].isnull().any():
            warnings.append("Alguns nomes estão vazios e serão ignorados")
        
        # Verificar nomes duplicados no arquivo
        duplicates = df[df.duplicated(subset=['nome_completo'], keep=False)]
        if not duplicates.empty:
            warnings.append(f"{len(duplicates)} nomes duplicados encontrados no arquivo")
        
        return {
            'valid': len(errors) == 0,
            'errors': errors,
            'warnings': warnings,
            'stats': {
                'total_rows': len(df),
                'valid_names': df['nome_completo'].notna().sum(),
                'duplicates': len(duplicates)
            }
        }

class ProgressCalculator:
    """
    Calculadora de progresso otimizada
    """
    
    @staticmethod
    def calculate_turma_progress(turma):
        """
        Calcula progresso de uma turma de forma otimizada
        """
        if not turma.competencias.exists():
            return {
                'progresso_percentual': 0,
                'alunos_completos': 0,
                'total_alunos': 0,
                'notas_lancadas': 0,
                'notas_possiveis': 0
            }
        
        alunos = turma.alunos.all()
        competencias = turma.competencias.all()
        
        total_alunos = len(alunos)
        total_competencias = len(competencias)
        total_notas_possiveis = total_alunos * total_competencias
        
        # Buscar todas as notas da turma de uma vez
        notas = LancamentoDeNota.objects.filter(
            aluno__in=alunos,
            competencia__in=competencias
        ).select_related('aluno', 'competencia')
        
        # Contar notas por aluno
        notas_por_aluno = {}
        for nota in notas:
            aluno_id = nota.aluno.pk
            if aluno_id not in notas_por_aluno:
                notas_por_aluno[aluno_id] = 0
            notas_por_aluno[aluno_id] += 1
        
        # Contar alunos completos
        alunos_completos = sum(
            1 for count in notas_por_aluno.values() 
            if count >= total_competencias
        )
        
        progresso_percentual = 0
        if total_notas_possiveis > 0:
            progresso_percentual = (len(notas) / total_notas_possiveis) * 100
        
        return {
            'progresso_percentual': int(progresso_percentual),
            'alunos_completos': alunos_completos,
            'total_alunos': total_alunos,
            'notas_lancadas': len(notas),
            'notas_possiveis': total_notas_possiveis
        }


class ProblemaSystemaManager:
    """
    Gerenciador para criação automática de problemas detectados pelo sistema
    """
    
    @staticmethod
    def criar_problema_automatico(tipo_problema, titulo, descricao, professor=None, turma=None, aluno=None, prioridade='MEDIA'):
        """
        Cria um problema detectado automaticamente pelo sistema
        """
        from .models import ProblemaRelatado
        from django.utils import timezone
        
        # Verifica se já existe um problema similar pendente
        problema_existente = ProblemaRelatado.objects.filter(
            origem='SISTEMA',
            tipo_problema=tipo_problema,
            status__in=['PENDENTE', 'EM_ANALISE'],
            turma=turma,
            aluno=aluno
        ).first()
        
        if problema_existente:
            # Atualiza a data do problema existente em vez de criar um novo
            problema_existente.data_atualizacao = timezone.now()
            problema_existente.save()
            return problema_existente
        
        # Cria novo problema
        problema = ProblemaRelatado.objects.create(
            origem='SISTEMA',
            professor=professor,
            turma=turma,
            aluno=aluno,
            tipo_problema=tipo_problema,
            titulo=titulo,
            descricao=descricao,
            prioridade=prioridade
        )
        
        logger.info(f"Problema automático criado: {problema}")
        return problema
    
    @staticmethod
    def detectar_e_criar_problemas_automaticos():
        """
        Detecta e cria problemas automaticamente baseado no estado do sistema
        """
        from .models import Turma, Aluno, Professor
        from django.db.models import Count
        
        problemas_criados = []
        
        # 1. Detectar turmas sem professor
        turmas_sem_professor = Turma.objects.filter(professor_responsavel__isnull=True)
        for turma in turmas_sem_professor:
            problema = ProblemaSystemaManager.criar_problema_automatico(
                tipo_problema='TURMA_SEM_PROFESSOR',
                titulo=f'Turma {turma.nome} sem professor responsável',
                descricao=f'A turma {turma.nome} ({turma.identificador_turma}) não possui um professor responsável atribuído.',
                turma=turma,
                prioridade='ALTA'
            )
            problemas_criados.append(problema)
        
        # 2. Detectar professores sem turma
        professores_sem_turma = Professor.objects.filter(turmas__isnull=True)
        for professor in professores_sem_turma:
            problema = ProblemaSystemaManager.criar_problema_automatico(
                tipo_problema='PROFESSOR_SEM_TURMA',
                titulo=f'Professor {professor.user.get_full_name() or professor.user.username} sem turma',
                descricao=f'O professor {professor.user.get_full_name() or professor.user.username} não possui turmas atribuídas.',
                professor=professor,
                prioridade='BAIXA'
            )
            problemas_criados.append(problema)
        
        # 3. Detectar alunos duplicados na mesma turma
        alunos_duplicados = Aluno.objects.values('nome_completo', 'turma').annotate(
            total=Count('id')
        ).filter(total__gt=1)
        
        for duplicado in alunos_duplicados:
            turma = Turma.objects.get(id=duplicado['turma'])
            alunos = Aluno.objects.filter(
                nome_completo=duplicado['nome_completo'],
                turma=turma
            )
            
            problema = ProblemaSystemaManager.criar_problema_automatico(
                tipo_problema='ALUNO_DUPLICADO',
                titulo=f'Aluno duplicado: {duplicado["nome_completo"]} na turma {turma.nome}',
                descricao=f'O aluno {duplicado["nome_completo"]} aparece {duplicado["total"]} vezes na turma {turma.nome}.',
                turma=turma,
                prioridade='ALTA'
            )
            problemas_criados.append(problema)
        
        return problemas_criados


class BoletimGenerator:
    """
    Classe para gerar boletins personalizados usando templates markdown com placeholders
    """
    
    # Mapeamento de tipos de boletim para arquivos de template
    TEMPLATE_MAP = {
        'adolescentes_adultos': 'adolescentes_adultos.docx',
        'material_antigo': 'material_antigo.docx',
        'lion_stars': 'lion_stars.docx',
        'junior': 'junior.docx',
    }
    
    # Mapeamento de competências para placeholders por tipo de boletim
    COMPETENCIA_PLACEHOLDERS = {
        'adolescentes_adultos': [
            'producao_oral',
            'producao_escrita',
            'avaliacoes_de_progresso',
            'nota_final',
        ],
        'material_antigo': [
            'producao_oral',
            'producao_escrita',
            'compreensao_oral',
            'compreensao_escrita',
            'writing_bit_01',
            'writing_bit_02',
            'checkpoints',
            'nota_final',
        ],
        'lion_stars': [
            'comunicacao_oral',
            'compreensao_oral',
            'interesse_pela_aprendizagem',
            'colaboracao',
            'engajamento',
        ],
        'junior': [
            'comunicacao_oral',
            'compreensao_oral',
            'comunicacao_escrita',
            'compreensao_escrita',
            'interesse_pela_aprendizagem',
            'colaboracao',
            'engajamento',
        ],
    }
    
    @staticmethod
    def _get_template_path(boletim_tipo):
        """
        Retorna o caminho completo do template de boletim
        """
        import os
        from django.conf import settings
        
        template_file = BoletimGenerator.TEMPLATE_MAP.get(boletim_tipo)
        if not template_file:
            raise ValueError(f"Tipo de boletim '{boletim_tipo}' não encontrado")
        
        # Caminho para a pasta de templates de boletim
        base_dir = settings.BASE_DIR
        template_path = os.path.join(base_dir, 'core', 'templates', 'boletins', template_file)
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template de boletim não encontrado: {template_path}")
        
        return template_path
    
    # Mapeamento de aliases de competências
    # Usado quando uma competência tem nomes diferentes mas representa a mesma coisa
    COMPETENCIA_ALIASES = {
        'compreensao_de_leitura': ['compreensao_escrita'],
        'compreensao_escrita': ['compreensao_de_leitura'],
    }
    
    @staticmethod
    def _normalizar_nome_competencia(nome_competencia):
        """
        Normaliza o nome da competência para corresponder aos placeholders
        Remove acentos, caracteres especiais e converte para snake_case
        """
        import unicodedata
        
        # Remove acentos
        nfkd = unicodedata.normalize('NFKD', nome_competencia)
        nome_sem_acento = "".join([c for c in nfkd if not unicodedata.combining(c)])
        
        # Converte para minúsculas e substitui espaços por underscore
        nome_normalizado = nome_sem_acento.lower()
        nome_normalizado = nome_normalizado.replace(' ', '_')
        
        # Remove caracteres especiais, mantendo apenas letras, números e underscore
        nome_normalizado = ''.join(c for c in nome_normalizado if c.isalnum() or c == '_')
        
        # Remove partes entre parênteses (como percentuais)
        if '(' in nome_normalizado:
            nome_normalizado = nome_normalizado.split('(')[0].strip('_')
        
        return nome_normalizado
    
    @staticmethod
    def _normalizar_texto(texto):
        """
        Remove acentos e caracteres especiais de um texto
        """
        import unicodedata
        nfkd = unicodedata.normalize('NFKD', texto)
        return "".join([c for c in nfkd if not unicodedata.combining(c)])
    
    @staticmethod
    def _substituir_em_paragraph(paragraph, substituicoes):
        """
        Substitui placeholders em um parágrafo, lidando com runs fragmentados e acentos
        """
        texto_original = paragraph.text
        texto_normalizado = BoletimGenerator._normalizar_texto(texto_original)
        substituiu_algo = False
        
        # Tenta encontrar placeholders normalizados
        for placeholder, valor in substituicoes.items():
            # Busca com e sem acentos
            placeholder_com_acentos_patterns = [
                f'<<{placeholder}>>',  # sem acentos (ex: <<nivel>>)
            ]
            
            # Gera variações com possíveis acentos
            if placeholder == 'nivel':
                placeholder_com_acentos_patterns.append('<<nível>>')
            elif placeholder == 'colaboracao':
                placeholder_com_acentos_patterns.append('<<colaboração>>')
            
            for pattern in placeholder_com_acentos_patterns:
                if pattern in texto_original:
                    print(f"  → Encontrado '{pattern}' no texto: '{texto_original[:50]}...'")
                    
                    # Substitui no texto completo
                    novo_texto = texto_original.replace(pattern, str(valor))
                    
                    # Limpa todos os runs
                    for run in paragraph.runs:
                        run.text = ''
                    
                    # Adiciona o texto substituído no primeiro run
                    if paragraph.runs:
                        paragraph.runs[0].text = novo_texto
                    else:
                        paragraph.add_run(novo_texto)
                    
                    print(f"  ✓ Substituído para: '{novo_texto[:50]}...'")
                    texto_original = novo_texto  # Atualiza para próximas substituições
                    substituiu_algo = True
                    break
        
        return substituiu_algo
    
    @staticmethod
    def gerar_boletim(aluno):
        """
        Gera o boletim personalizado para um aluno específico a partir de template Word (.docx)
        
        Args:
            aluno: Instância do model Aluno
            
        Returns:
            Document: Objeto Document do python-docx com placeholders preenchidos
            
        Raises:
            ValueError: Se o tipo de boletim da turma não for válido
            FileNotFoundError: Se o template não for encontrado
        """
        from docx import Document
        
        turma = aluno.turma
        boletim_tipo = turma.boletim_tipo
        
        # Carrega o template Word
        template_path = BoletimGenerator._get_template_path(boletim_tipo)
        doc = Document(template_path)
        
        # Dados básicos do aluno
        professor_nome = 'N/A'
        if turma.professor_responsavel:
            # Prioriza nome_completo, depois full_name do User, depois username
            if turma.professor_responsavel.nome_completo:
                professor_nome = turma.professor_responsavel.nome_completo
            else:
                professor_nome = turma.professor_responsavel.user.get_full_name() or turma.professor_responsavel.user.username
        
        dados_basicos = {
            'aluno': aluno.nome_completo,
            'nivel': turma.tipo_turma.nome if turma.tipo_turma else turma.identificador_turma,
            'professor': professor_nome,
        }
        
        print(f"\n{'='*80}")
        print(f"DEBUG - Dados básicos:")
        print(f"  aluno: {dados_basicos['aluno']}")
        print(f"  nivel: {dados_basicos['nivel']}")
        print(f"  professor: {dados_basicos['professor']}")
        print(f"{'='*80}\n")
        
        # Busca as notas do aluno (retorna uma lista)
        notas_list = aluno.get_notas_boletim()
        
        # Preparar todos os placeholders para substituição
        placeholders_esperados = BoletimGenerator.COMPETENCIA_PLACEHOLDERS.get(boletim_tipo, [])
        
        print(f"Boletim tipo: {boletim_tipo}")
        print(f"Placeholders esperados: {placeholders_esperados}")
        print(f"Total de notas: {len(notas_list)}")
        
        # Dicionário completo de substituições
        substituicoes = dados_basicos.copy()
        
        # Adicionar notas ao dicionário de substituições
        print("\nBUSCANDO NOTAS:")
        for placeholder in placeholders_esperados:
            nota_encontrada = 'N/A'
            print(f"\n  Procurando por placeholder: '{placeholder}'")
            
            # Lista de nomes a buscar (placeholder + seus aliases)
            nomes_busca = [placeholder]
            if placeholder in BoletimGenerator.COMPETENCIA_ALIASES:
                nomes_busca.extend(BoletimGenerator.COMPETENCIA_ALIASES[placeholder])
                print(f"    (Aliases: {BoletimGenerator.COMPETENCIA_ALIASES[placeholder]})")
            
            # notas_list é uma lista de dicionários com 'competencia', 'nota', 'data_lancamento'
            for nota_info in notas_list:
                competencia = nota_info.get('competencia')
                if competencia:
                    competencia_normalizada = BoletimGenerator._normalizar_nome_competencia(competencia.nome)
                    print(f"    Comparando: '{competencia.nome}' -> '{competencia_normalizada}'")
                    
                    # Verifica match exato com o placeholder ou qualquer um dos aliases
                    if competencia_normalizada in nomes_busca:
                        nota_valor = nota_info.get('nota', 'N/A')
                        nota_encontrada = str(nota_valor) if nota_valor != '-' else 'N/A'
                        print(f"    ✓ MATCH EXATO! Nota: {nota_encontrada}")
                        break
                    # Verifica match parcial apenas se não houver match exato
                    elif any(nome in competencia_normalizada for nome in nomes_busca):
                        nota_valor = nota_info.get('nota', 'N/A')
                        nota_encontrada = str(nota_valor) if nota_valor != '-' else 'N/A'
                        print(f"    ✓ MATCH PARCIAL! Nota: {nota_encontrada}")
                        break
            
            if nota_encontrada == 'N/A':
                print(f"    ✗ Nenhuma correspondência encontrada")
            
            substituicoes[placeholder] = nota_encontrada
        
        # Calcular nota_final se for material_antigo ou adolescentes_adultos
        if boletim_tipo in ['material_antigo', 'adolescentes_adultos']:
            # Mapeamento de conceitos para valores numéricos para calcular média
            conceito_para_numero = {
                'A': 95,  # 100-90%
                'B': 82,  # 89-75%
                'C': 67,  # 74-60%
                'D': 50,  # 59% ou menos
            }
            
            # Para adolescentes_adultos: producao_oral (40%) + producao_escrita (40%) + avaliacoes_de_progresso (20%)
            if boletim_tipo == 'adolescentes_adultos':
                nota_producao_oral = None
                nota_producao_escrita = None
                nota_avaliacoes = None
                
                for nota_info in notas_list:
                    competencia = nota_info.get('competencia')
                    nota_valor = nota_info.get('nota')
                    
                    if competencia and nota_valor and nota_valor != '-':
                        competencia_norm = BoletimGenerator._normalizar_nome_competencia(competencia.nome)
                        
                        # Converte conceito para número
                        valor_numerico = None
                        if str(nota_valor).upper() in conceito_para_numero:
                            valor_numerico = conceito_para_numero[str(nota_valor).upper()]
                        else:
                            try:
                                valor_numerico = float(nota_valor)
                            except (ValueError, TypeError):
                                pass
                        
                        if valor_numerico is not None:
                            if competencia_norm == 'producao_oral':
                                nota_producao_oral = valor_numerico
                            elif competencia_norm == 'producao_escrita':
                                nota_producao_escrita = valor_numerico
                            elif competencia_norm == 'avaliacoes_de_progresso':
                                nota_avaliacoes = valor_numerico
                
                # Calcula média ponderada: 40% + 40% + 20%
                if nota_producao_oral is not None and nota_producao_escrita is not None and nota_avaliacoes is not None:
                    media = (nota_producao_oral * 0.4) + (nota_producao_escrita * 0.4) + (nota_avaliacoes * 0.2)
                    print(f"\n  📊 Cálculo ponderado:")
                    print(f"      Produção Oral (40%): {nota_producao_oral:.1f}")
                    print(f"      Produção Escrita (40%): {nota_producao_escrita:.1f}")
                    print(f"      Avaliações de Progresso (20%): {nota_avaliacoes:.1f}")
                    print(f"      Média ponderada: {media:.1f}")
                elif nota_producao_oral is not None and nota_producao_escrita is not None:
                    # Se não tem avaliacoes, faz média das duas (50% cada)
                    media = (nota_producao_oral + nota_producao_escrita) / 2
                    print(f"\n  📊 Cálculo sem avaliações (média simples de 2 notas): {media:.1f}")
                else:
                    media = None
                    print(f"\n  ⚠️ Notas insuficientes para calcular média")
                
                if media is not None:
                    # Converte média numérica de volta para conceito
                    if media >= 90:
                        nota_final_conceito = 'A'
                    elif media >= 75:
                        nota_final_conceito = 'B'
                    elif media >= 60:
                        nota_final_conceito = 'C'
                    else:
                        nota_final_conceito = 'D'
                    
                    substituicoes['nota_final'] = nota_final_conceito
                    print(f"      Nota Final: {nota_final_conceito}")
                else:
                    substituicoes['nota_final'] = 'N/A'
            
            # Para material_antigo: média simples de todas as notas
            else:
                valores_para_media = []
                
                for nota_info in notas_list:
                    nota_valor = nota_info.get('nota')
                    if nota_valor and nota_valor != '-':
                        # Converte conceito para número
                        if str(nota_valor).upper() in conceito_para_numero:
                            valores_para_media.append(conceito_para_numero[str(nota_valor).upper()])
                        else:
                            # Se for numérico, usa direto
                            try:
                                nota_num = float(nota_valor)
                                valores_para_media.append(nota_num)
                            except (ValueError, TypeError):
                                pass
                
                # Calcula a média e converte de volta para conceito
                if valores_para_media:
                    media = sum(valores_para_media) / len(valores_para_media)
                    
                    # Converte média numérica de volta para conceito
                    if media >= 90:
                        nota_final_conceito = 'A'
                    elif media >= 75:
                        nota_final_conceito = 'B'
                    elif media >= 60:
                        nota_final_conceito = 'C'
                    else:
                        nota_final_conceito = 'D'
                    
                    substituicoes['nota_final'] = nota_final_conceito
                    print(f"\n  📊 Nota Final calculada: {nota_final_conceito} (média {media:.1f} de {len(valores_para_media)} notas)")
                else:
                    substituicoes['nota_final'] = 'N/A'
                    print(f"\n  ⚠️ Nota Final: N/A (sem notas válidas)")
        
        print(f"\nDicionário de substituições completo:")
        for key, value in substituicoes.items():
            print(f"  <<{key}>> → {value}")
        print()
        
        # DEBUG: Mostrar todos os textos do documento
        print("="*80)
        print("TEXTOS ENCONTRADOS NO DOCUMENTO:")
        print("="*80)
        print("\nParágrafos:")
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                print(f"  [{i}] {paragraph.text[:100]}")
        
        print("\nTabelas:")
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    for pi, paragraph in enumerate(cell.paragraphs):
                        if paragraph.text.strip():
                            print(f"  [T{ti}-R{ri}-C{ci}-P{pi}] {paragraph.text[:100]}")
        print("="*80)
        print()
        
        # Substituir placeholders no documento Word
        substituicoes_feitas = []
        
        print("INICIANDO SUBSTITUIÇÕES:")
        # Substituir em parágrafos
        for paragraph in doc.paragraphs:
            if BoletimGenerator._substituir_em_paragraph(paragraph, substituicoes):
                substituicoes_feitas.append(f"Parágrafo substituído")
        
        # Substituir em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if BoletimGenerator._substituir_em_paragraph(paragraph, substituicoes):
                            substituicoes_feitas.append(f"Célula de tabela substituída")
        
        print(f"Total de parágrafos/células processados: {len(substituicoes_feitas)}")
        print()
        
        logger.info(f"Boletim gerado com sucesso para aluno {aluno.nome_completo} (Turma: {turma.nome})")
        
        return doc
    
    @staticmethod
    def gerar_boletim_turma(turma):
        """
        Gera boletims para todos os alunos de uma turma
        
        Args:
            turma: Instância do model Turma
            
        Returns:
            dict: Dicionário com {aluno_id: Document object}
        """
        boletins = {}
        alunos = turma.alunos.filter(ativo=True)
        
        for aluno in alunos:
            try:
                boletim_doc = BoletimGenerator.gerar_boletim(aluno)
                boletins[aluno.id] = boletim_doc
            except Exception as e:
                logger.error(f"Erro ao gerar boletim para aluno {aluno.nome_completo}: {str(e)}")
                boletins[aluno.id] = None
        
        return boletins